# contract_generator.py
# Генерация договора купли-продажи в рассрочку

import io
import copy
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Cm, Inches
from num2words import num2words


def number_to_words_ru(n: int) -> str:
    """Переводит целое число в русские слова"""
    try:
        return num2words(int(n), lang='ru')
    except:
        return str(n)


def money_to_words(amount: float) -> str:
    """Сумма прописью + 'сом'"""
    amount = int(round(amount))
    words = number_to_words_ru(amount)
    # Согласование "сом"
    last = amount % 10
    last2 = amount % 100
    if last2 >= 11 and last2 <= 19:
        som = "сом"
    elif last == 1:
        som = "сом"
    elif last in (2, 3, 4):
        som = "сома"
    else:
        som = "сом"
    return f"{words} {som}"


def months_to_words(months: int) -> str:
    """Количество месяцев прописью"""
    return number_to_words_ru(months)


def generate_payment_schedule(total_amount: float, down_payment: float, months: int, start_date: datetime = None):
    """
    Строит равномерный график платежей.
    total_amount — полная стоимость товара
    down_payment — первоначальный взнос
    months — срок рассрочки
    """
    if start_date is None:
        start_date = datetime.now()

    remaining = max(0.0, float(total_amount) - float(down_payment))
    if months <= 0:
        months = 1

    monthly = round(remaining / months, 2)
    schedule = []
    balance = remaining

    for i in range(1, months + 1):
        pay_date = start_date + timedelta(days=30 * i)
        if i == months:
            amount = round(balance, 2)
            balance = 0.0
        else:
            amount = monthly
            balance = round(balance - monthly, 2)

        schedule.append({
            "num": i,
            "date": pay_date.strftime("%d.%m.%Y"),
            "amount": amount,
            "balance": max(0.0, balance)
        })

    return schedule


def fill_contract(
    template_path: str,
    contract_number: str,
    contract_date: str,
    client_name: str,
    client_address: str,
    client_passport: str,
    total_amount: float,
    months: int,
    product_name: str,
    product_qty: int,
    product_price: float,
    down_payment: float = 0,
    schedule: list = None,
) -> bytes:
    """
    Заполняет шаблон договора и возвращает байты готового .docx
    """
    doc = Document(template_path)

    # Готовим значения
    total_amount = float(total_amount)
    months = int(months)
    product_qty = int(product_qty)
    product_price = float(product_price)
    down_payment = float(down_payment or 0)

    amount_words = money_to_words(total_amount)
    months_words = months_to_words(months)

    if schedule is None:
        schedule = generate_payment_schedule(total_amount, down_payment, months)

    # Словарь замен
    replacements = {
        "{{CONTRACT_NUMBER}}": str(contract_number),
        "{{CONTRACT_DATE}}": contract_date,
        "{{CLIENT_NAME}}": client_name or "",
        "{{CLIENT_ADDRESS}}": client_address or "—",
        "{{CLIENT_PASSPORT}}": client_passport or "—",
        "{{TOTAL_AMOUNT}}": f"{total_amount:,.0f}".replace(",", " "),
        "{{TOTAL_AMOUNT_WORDS}}": amount_words,
        "{{MONTHS}}": str(months),
        "{{MONTHS_WORDS}}": months_words,
        "{{PRODUCT_NAME}}": product_name or "",
        "{{PRODUCT_QTY}}": str(product_qty),
        "{{PRODUCT_PRICE}}": f"{product_price:,.0f}".replace(",", " "),
        "{{SCHEDULE_TOTAL}}": f"{sum(s['amount'] for s in schedule):,.0f}".replace(",", " "),
    }

    # Заполняем график (до 12 месяцев)
    for i in range(1, 13):
        if i <= len(schedule):
            s = schedule[i - 1]
            replacements[f"{{{{P{i}_NUM}}}}"] = str(s["num"])
            replacements[f"{{{{P{i}_DATE}}}}"] = s["date"]
            replacements[f"{{{{P{i}_AMOUNT}}}}"] = f"{s['amount']:,.0f}".replace(",", " ")
            replacements[f"{{{{P{i}_BALANCE}}}}"] = f"{s['balance']:,.0f}".replace(",", " ")
        else:
            replacements[f"{{{{P{i}_NUM}}}}"] = ""
            replacements[f"{{{{P{i}_DATE}}}}"] = ""
            replacements[f"{{{{P{i}_AMOUNT}}}}"] = ""
            replacements[f"{{{{P{i}_BALANCE}}}}"] = ""

    def replace_in_paragraph(paragraph):
        full_text = paragraph.text
        if not full_text:
            return
        for key, value in replacements.items():
            if key in full_text:
                full_text = full_text.replace(key, str(value))
        # Перезаписываем runs
        if paragraph.runs:
            # Сохраняем стиль первого run
            first_run = paragraph.runs[0]
            for run in paragraph.runs:
                run.text = ""
            first_run.text = full_text

    # Замена во всех параграфах
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    # Замена в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)

    # Сохраняем в память
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ========== Тест ==========
if __name__ == "__main__":
    schedule = generate_payment_schedule(
        total_amount=109000,
        down_payment=9000,
        months=10
    )
    print("График платежей:")
    for s in schedule:
        print(s)

    data = fill_contract(
        template_path="/home/workdir/artifacts/contract_template.docx",
        contract_number="15",
        contract_date="24.07.2026",
        client_name="Абдимоминов Мурод",
        client_address="г. Ош, ул. Ленина 12",
        client_passport="AN1234567 от 15.03.2020",
        total_amount=109000,
        months=10,
        product_name="Стиральная машина LG",
        product_qty=1,
        product_price=109000,
        down_payment=9000,
        schedule=schedule,
    )
    with open("/home/workdir/artifacts/test_contract_filled.docx", "wb") as f:
        f.write(data)
    print("\nТестовый договор сохранён: test_contract_filled.docx")
